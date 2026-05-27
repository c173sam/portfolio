"""
简历生成器 — 根据原始内容创建排版精美的 Word 简历
可方便的根据岗位 JD 修改各模块内容
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 📝 在这里修改你的简历内容 — 按 JD 自由调整
# ============================================================

# --- 基本信息 ---
INFO = {
    "name": "徐念齐",
    "title": "IT 网络运维工程师",
    "phone": "138-xxxx-xxxx",
    "email": "nianqixu@example.com",
    "location": "上海 / 深圳",
    "age": "XX岁",
}

# --- 个人概述（根据目标岗位调整关键词） ---
SUMMARY = (
    "拥有 2 年 IT 运维实战经验，熟悉服务器日常巡检、硬件故障诊断及网络安全防护。"
    "擅长跨部门沟通与现场应急响应，能快速定位并解决设备故障，保障业务连续性。"
    "注重预防性维护与流程优化，通过数据驱动的运维策略降低非计划停机时间。"
)

# --- 核心技能（按 JD 勾选或增删） ---
SKILLS = [
    ("服务器运维", "Linux/Windows 服务器巡检、性能监控（CPU/内存/磁盘 I/O）、系统部署"),
    ("网络管理", "TCP/IP 协议栈、VLAN 划分、网络设备（交换机/路由器）基础配置"),
    ("硬件维护", "服务器硬件诊断、易损件更换、PDA/自动化设备现场维修"),
    ("安全防护", "木马/病毒查杀、防火墙策略配置、白名单管理、漏洞修复"),
    ("监控工具", "Zabbix / Prometheus / Grafana 基础使用，日志分析与告警配置"),
    ("脚本能力", "Python / Shell 脚本编写，运维自动化任务开发"),
    ("文档协作", "运维文档撰写、知识库维护、跨部门技术对接"),
]

# --- 工作经历（按时间倒序，每条经验可独立增删） ---
EXPERIENCES = [
    {
        "company": "深圳纹声动力公司",
        "role": "采集机运维工程师",
        "period": "2024.03 — 至今",
        "bullets": [
            "独立负责采集机硬件状态检视与易损部件更换，确保设备在高压环境下稳定运行，未发生因硬件故障导致采集任务中断",
            "在维护过程中主动发现并清除潜伏的木马病毒，通过更新白名单策略强化前端设备安全基线，保障数据源头安全",
            "记录每日设备运行数据，总结易损部件损耗规律，协助团队制定更合理的备件采购计划，降低非计划停机时间",
        ],
    },
    {
        "company": "上海德邦物流公司",
        "role": "服务器运维（实习）",
        "period": "2023.06 — 2024.01",
        "bullets": [
            "负责物流园区服务器及网络设备日常巡检，监控 CPU、内存、磁盘 I/O 等性能指标，确保分拣系统、仓储系统等核心业务稳定运行",
            "快速响应现场设备故障（手持终端 PDA、自动化分拣设备），完成硬件排查与易损件更换，保障一线操作人员作业不中断，实习期间协助处理硬件故障 30+ 起",
            "巡检过程中发现并清除潜伏的木马病毒，通过隔离受感染终端、修复漏洞及优化安全策略，防止病毒在内网横向传播",
            "担任现场技术接口人，协调 IT 团队与现场操作人员工作安排，在系统升级或设备维护期间合理规划停机时间，避免影响物流高峰时段",
        ],
    },
]

# --- 教育背景 ---
EDUCATION = {
    "school": "山东轻工职业学院",
    "major": "现代通信技术",
    "degree": "大专",
    "period": "2021.09 — 2024.06",
    "courses": "通信原理、网络基础、Linux 系统管理、设备运维",
}

# --- 证书与资质（按需填写） ---
CERTIFICATIONS = [
    "网络管理员（中级）",
    "电工证（低压维修）",
    # "RHCSA",  # 如果有可以加上
    # "CCNA",
]

# --- 自我评价（根据 JD 调整侧重点） ---
SELF_ASSESSMENT = [
    "拥有服务器运维及现场设备维护经验，熟悉服务器与采集设备日常巡检、硬件故障诊断及病毒查杀，保障系统稳定运行",
    "善于跨部门沟通与现场人员协调，应急响应冷静高效，有效降低业务中断风险",
    "注重流程优化，通过预防性维护与总结改进，提升团队整体运维效率",
]

# ============================================================
# 🔧 排版引擎 — 以下为自动排版代码，一般无需修改
# ============================================================


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_horizontal_line(doc, color='1a1a2e', width=Pt(1)):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_title(doc, title, color='1a1a2e'):
    """添加带底色的章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    # 添加字符底纹效果
    run = p.add_run(f'  {title}  ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 底纹通过 shading 实现
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    rPr.append(shd)


def add_body_text(doc, text, indent=False, size=Pt(10), bold=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text, size=Pt(9.5)):
    """添加项目符号"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run('▸ ' + text)
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def build_resume():
    doc = Document()

    # --- 全局样式 ---
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # --- 页面设置 ---
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ==================== 头部区域 ====================
    # 姓名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(INFO['name'])
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # 目标岗位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(INFO['title'])
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    run.bold = True

    # 联系方式 — 用表格实现居中分布
    contact_text = f"{INFO['phone']}  |  {INFO['email']}  |  {INFO['location']}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(contact_text)
    run.font.size = Pt(9)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x88)

    add_horizontal_line(doc, color='1a1a2e')

    # ==================== 个人概述 ====================
    add_section_title(doc, '个人概述')
    add_body_text(doc, SUMMARY, indent=True)

    # ==================== 核心技能 ====================
    add_section_title(doc, '核心技能')

    # 将技能排成三列
    skills_per_row = 3
    for i in range(0, len(SKILLS), skills_per_row):
        row_skills = SKILLS[i:i+skills_per_row]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.5
        for j, (name, desc) in enumerate(row_skills):
            if j > 0:
                run = p.add_run('    ')
                run.font.size = Pt(9)
            run = p.add_run(f'{name}：')
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            run = p.add_run(desc)
            run.font.size = Pt(8.5)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ==================== 工作经历 ====================
    add_section_title(doc, '工作经历')

    for exp in EXPERIENCES:
        # 公司 - 职位 - 时间
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{exp['company']}  ")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        run = p.add_run(f"{exp['role']}")
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

        run = p.add_run(f"    {exp['period']}")
        run.font.size = Pt(8.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        for bullet in exp['bullets']:
            add_bullet(doc, bullet)

    # ==================== 教育背景 ====================
    add_section_title(doc, '教育背景')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"{EDUCATION['school']}  ")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    run = p.add_run(f"{EDUCATION['major']}（{EDUCATION['degree']}）    {EDUCATION['period']}")
    run.font.size = Pt(9.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    add_body_text(doc, f"主修课程：{EDUCATION['courses']}", indent=True, size=Pt(9))

    # ==================== 证书与资质 ====================
    if CERTIFICATIONS:
        add_section_title(doc, '证书与资质')
        cert_text = '    '.join([f'〔{c}〕' for c in CERTIFICATIONS])
        add_body_text(doc, cert_text, indent=True)

    # ==================== 自我评价 ====================
    add_section_title(doc, '自我评价')
    for item in SELF_ASSESSMENT:
        add_bullet(doc, item)

    # ==================== 页脚声明 ====================
    add_horizontal_line(doc, color='cccccc')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('本人承诺以上信息真实有效，期待与贵司共成长。')
    run.font.size = Pt(8)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    # --- 保存 ---
    output_path = r'C:\Users\ROG\Desktop\徐念齐-简历-优化版.docx'
    doc.save(output_path)
    print(f'✅ 简历已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    build_resume()
